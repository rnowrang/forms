"""Template service for DOCX parsing and schema extraction."""

import os
import re
import shutil
from typing import Optional, List, Dict, Any
from datetime import datetime

from fastapi import UploadFile, HTTPException, status
from sqlalchemy.orm import Session
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.config import get_settings
from app.models.template import Template
from app.schemas.template import TemplateCreate, TemplateUpdate, TemplateSchema

settings = get_settings()


class TemplateService:
    """Service for template management and DOCX processing."""
    
    @staticmethod
    def get_template(db: Session, template_id: int) -> Optional[Template]:
        """Get a template by ID."""
        return db.query(Template).filter(Template.id == template_id).first()
    
    @staticmethod
    def get_templates(
        db: Session, 
        skip: int = 0, 
        limit: int = 100,
        active_only: bool = True
    ) -> List[Template]:
        """Get all templates."""
        query = db.query(Template)
        if active_only:
            query = query.filter(Template.is_active == True)
        return query.offset(skip).limit(limit).all()
    
    @staticmethod
    def get_published_templates(db: Session) -> List[Template]:
        """Get all published templates."""
        return db.query(Template).filter(
            Template.is_active == True,
            Template.is_published == True
        ).all()
    
    @staticmethod
    async def create_template(
        db: Session,
        template_data: TemplateCreate,
        file: UploadFile
    ) -> Template:
        """Create a new template from an uploaded DOCX file."""
        # Validate file type
        if not file.filename.endswith(('.docx', '.doc')):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File must be a DOCX or DOC file"
            )
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename)
        stored_filename = f"{timestamp}_{safe_filename}"
        file_path = os.path.join(settings.template_dir, stored_filename)
        
        # Save file
        os.makedirs(settings.template_dir, exist_ok=True)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract schema from DOCX
        try:
            schema = TemplateService.extract_schema_from_docx(file_path)
        except Exception as e:
            # Clean up file on error
            os.remove(file_path)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to parse DOCX: {str(e)}"
            )
        
        # Create template record
        db_template = Template(
            name=template_data.name,
            description=template_data.description,
            version=template_data.version,
            original_file_path=file_path,
            original_file_name=file.filename,
            schema=schema,
        )
        db.add(db_template)
        db.commit()
        db.refresh(db_template)
        
        return db_template
    
    @staticmethod
    def update_template(
        db: Session,
        template_id: int,
        template_data: TemplateUpdate
    ) -> Optional[Template]:
        """Update a template."""
        db_template = TemplateService.get_template(db, template_id)
        if not db_template:
            return None
        
        update_data = template_data.model_dump(exclude_unset=True)
        
        # Handle schema update specially
        if "schema" in update_data and update_data["schema"]:
            update_data["schema"] = update_data["schema"].model_dump() if hasattr(update_data["schema"], "model_dump") else update_data["schema"]
        
        for key, value in update_data.items():
            setattr(db_template, key, value)
        
        db.commit()
        db.refresh(db_template)
        return db_template
    
    @staticmethod
    def delete_template(db: Session, template_id: int) -> bool:
        """Soft delete a template (mark as inactive)."""
        db_template = TemplateService.get_template(db, template_id)
        if not db_template:
            return False
        
        db_template.is_active = False
        db.commit()
        return True
    
    @staticmethod
    def extract_schema_from_docx(file_path: str) -> Dict[str, Any]:
        """
        Extract form schema from a DOCX file.
        
        This analyzes the document structure to identify:
        - Sections (Roman numerals, headers)
        - Fields (text inputs, checkboxes, tables with input cells)
        - Anchors (how to locate each field in the document)
        """
        doc = Document(file_path)
        
        sections = []
        fields = []
        current_section = None
        section_order = 0
        field_order = 0
        
        # Roman numeral pattern for sections
        roman_pattern = re.compile(r'^([IVXLC]+)\.\s+(.+)$', re.IGNORECASE)
        
        # Pattern for field labels (text followed by colon or underscores)
        field_label_pattern = re.compile(r'^(.+?):\s*(_+|$)')
        checkbox_pattern = re.compile(r'[☐☑✓□■◻◼]')
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Check for section headers
            roman_match = roman_pattern.match(text)
            if roman_match:
                section_id = f"sec_{roman_match.group(1).upper()}"
                current_section = {
                    "id": section_id,
                    "title": text,
                    "description": None,
                    "order": section_order,
                    "collapsible": True,
                    "collapsed_by_default": False,
                }
                sections.append(current_section)
                section_order += 1
                continue
            
            # Check for field labels
            field_match = field_label_pattern.match(text)
            if field_match and current_section:
                label = field_match.group(1).strip()
                field_id = TemplateService._generate_field_id(label, current_section["id"])
                
                field = {
                    "id": field_id,
                    "type": "text",
                    "label": label,
                    "section_id": current_section["id"],
                    "required": False,
                    "anchor": {
                        "type": "paragraph",
                        "paragraph_contains": label,
                        "paragraph_index": para_idx,
                    },
                    "order": field_order,
                }
                fields.append(field)
                field_order += 1
            
            # Check for checkboxes
            if checkbox_pattern.search(text) and current_section:
                # Extract checkbox options
                options = TemplateService._extract_checkbox_options(text)
                if options:
                    label = text.split(options[0]["label"])[0].strip().rstrip(':')
                    if not label:
                        label = "Options"
                    
                    field_id = TemplateService._generate_field_id(label, current_section["id"])
                    field = {
                        "id": field_id,
                        "type": "checkbox",
                        "label": label,
                        "section_id": current_section["id"],
                        "required": False,
                        "options": options,
                        "anchor": {
                            "type": "paragraph",
                            "paragraph_contains": text[:50],
                            "paragraph_index": para_idx,
                        },
                        "order": field_order,
                    }
                    fields.append(field)
                    field_order += 1
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_fields = TemplateService._extract_table_fields(
                table, table_idx, current_section, field_order
            )
            fields.extend(table_fields)
            field_order += len(table_fields)
        
        # Build default rules based on common patterns
        rules = TemplateService._generate_default_rules(fields)
        
        return {
            "sections": sections,
            "fields": fields,
            "rules": rules,
        }
    
    @staticmethod
    def _generate_field_id(label: str, section_id: str) -> str:
        """Generate a stable field ID from label."""
        # Clean and normalize the label
        clean = re.sub(r'[^a-zA-Z0-9\s]', '', label.lower())
        clean = re.sub(r'\s+', '_', clean.strip())
        # Truncate if too long
        if len(clean) > 50:
            clean = clean[:50]
        return f"{section_id}.{clean}"
    
    @staticmethod
    def _extract_checkbox_options(text: str) -> List[Dict[str, str]]:
        """Extract checkbox options from text."""
        options = []
        # Split by checkbox characters
        checkbox_chars = r'[☐☑✓□■◻◼]'
        parts = re.split(checkbox_chars, text)
        
        for i, part in enumerate(parts[1:], 1):  # Skip first part (before first checkbox)
            label = part.strip()
            if label:
                # Clean up the label
                label = re.sub(r'\s+', ' ', label)
                options.append({
                    "value": f"option_{i}",
                    "label": label.split('\n')[0].strip()  # First line only
                })
        
        return options
    
    @staticmethod
    def _extract_table_fields(
        table: Table,
        table_idx: int,
        current_section: Optional[Dict],
        start_order: int
    ) -> List[Dict[str, Any]]:
        """Extract fields from a table."""
        fields = []
        section_id = current_section["id"] if current_section else "sec_general"
        
        if len(table.rows) < 2:
            return fields
        
        # Try to identify header row
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        
        # Check if this is a data entry table (has recognizable headers)
        data_headers = ["name", "title", "role", "date", "email", "phone"]
        is_data_table = any(
            any(h.lower() in header.lower() for h in data_headers)
            for header in headers
        )
        
        if is_data_table:
            # Create a repeatable field group
            field_id = f"{section_id}.table_{table_idx}"
            field = {
                "id": field_id,
                "type": "repeatable",
                "label": "Table Data",
                "section_id": section_id,
                "required": False,
                "repeatable_config": {
                    "min_rows": 1,
                    "max_rows": 20,
                    "columns": [
                        {"id": f"col_{i}", "label": h, "type": "text"}
                        for i, h in enumerate(headers) if h
                    ],
                },
                "anchor": {
                    "type": "table",
                    "table_index": table_idx,
                    "header_text": headers[0] if headers else None,
                },
                "order": start_order,
            }
            fields.append(field)
        else:
            # Process as label-value pairs
            for row_idx, row in enumerate(table.rows):
                cells = row.cells
                if len(cells) >= 2:
                    label = cells[0].text.strip()
                    if label and label.endswith(':'):
                        label = label[:-1]
                    
                    if label:
                        field_id = TemplateService._generate_field_id(label, section_id)
                        field = {
                            "id": field_id,
                            "type": "text",
                            "label": label,
                            "section_id": section_id,
                            "required": False,
                            "anchor": {
                                "type": "table_cell",
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "column_index": 1,
                                "label_text": label,
                            },
                            "order": start_order + row_idx,
                        }
                        fields.append(field)
        
        return fields
    
    @staticmethod
    def _generate_default_rules(fields: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Generate default conditional rules based on common patterns."""
        rules = []
        
        # Look for "Other" options in checkbox/select fields
        for field in fields:
            if field["type"] in ["checkbox", "select", "radio"]:
                options = field.get("options", [])
                has_other = any(
                    "other" in opt.get("label", "").lower()
                    for opt in options
                )
                
                if has_other:
                    # Check if there's a corresponding "other specify" field
                    other_field_id = f"{field['id']}_other"
                    other_field_exists = any(
                        f["id"] == other_field_id or "other" in f["id"].lower()
                        for f in fields
                    )
                    
                    if not other_field_exists:
                        # Create rule to show/hide other specify
                        rules.append({
                            "id": f"rule_{field['id']}_other",
                            "conditions": [
                                {
                                    "field": field["id"],
                                    "operator": "contains",
                                    "value": "other"
                                }
                            ],
                            "then_actions": [
                                {"action": "show", "field": f"{field['id']}_specify"}
                            ],
                            "else_actions": [
                                {"action": "hide", "field": f"{field['id']}_specify"},
                                {"action": "clear", "field": f"{field['id']}_specify"}
                            ],
                        })
        
        return rules
