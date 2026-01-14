"""Document generation service for DOCX and PDF output."""

import os
import shutil
import subprocess
import re
from typing import Optional, Dict, Any, Tuple
from datetime import datetime

from sqlalchemy.orm import Session
from fastapi import HTTPException, status
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

from app.config import get_settings
from app.models.template import Template
from app.models.form import FormInstance, FormVersion, FormData
from docx.oxml.ns import qn
from lxml import etree

settings = get_settings()


class DocumentService:
    """Service for generating filled DOCX and PDF documents."""

    @staticmethod
    def _check_legacy_checkbox(para, check: bool = True) -> bool:
        """
        Check or uncheck a legacy FORMCHECKBOX in a paragraph.

        Legacy checkboxes have structure:
        <w:fldChar w:fldCharType="begin">
          <w:ffData>
            <w:checkBox>
              <w:checked w:val="1"/>  <!-- Add this to check -->
            </w:checkBox>
          </w:ffData>
        </w:fldChar>

        Returns True if checkbox was found and updated.
        """
        for run in para.runs:
            # Find ffData element containing checkbox
            ffData = run._r.find('.//' + qn('w:ffData'))
            if ffData is not None:
                checkbox = ffData.find(qn('w:checkBox'))
                if checkbox is not None:
                    # Remove existing checked element if present
                    existing_checked = checkbox.find(qn('w:checked'))
                    if existing_checked is not None:
                        checkbox.remove(existing_checked)

                    # Add checked element if checking
                    if check:
                        checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                        checked_elem.set(qn('w:val'), '1')

                    return True
        return False

    @staticmethod
    def _check_checkbox_by_text(doc, search_text: str, check: bool = True) -> bool:
        """
        Find a paragraph containing search_text and check its FORMCHECKBOX.

        Returns True if checkbox was found and checked.
        """
        search_lower = search_text.lower()
        for para in doc.paragraphs:
            if search_lower in para.text.lower():
                if DocumentService._check_legacy_checkbox(para, check):
                    return True
        return False

    @staticmethod
    def _check_checkbox_in_table(doc, table_index: int, search_text: str, check: bool = True) -> bool:
        """
        Find a cell in a table containing search_text and check its FORMCHECKBOX.

        Returns True if checkbox was found and checked.
        """
        if table_index >= len(doc.tables):
            return False

        table = doc.tables[table_index]
        search_lower = search_text.lower()

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if search_lower in para.text.lower():
                        if DocumentService._check_legacy_checkbox(para, check):
                            return True
        return False

    @staticmethod
    def _flatten_dict(d: Dict[str, Any], parent_key: str = '', sep: str = '.') -> Dict[str, Any]:
        """
        Flatten a nested dictionary to dot-notation keys.

        Example: {'personnel': {'pi_name': 'John'}} -> {'personnel.pi_name': 'John'}
        """
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(DocumentService._flatten_dict(v, new_key, sep).items())
            else:
                items.append((new_key, v))
        return dict(items)

    @staticmethod
    def generate_documents(
        db: Session,
        form_id: int,
        version_id: Optional[int] = None
    ) -> Tuple[str, str]:
        """
        Generate filled DOCX and PDF for a form version.
        
        Returns tuple of (docx_path, pdf_path).
        """
        # Get form and template
        form = db.query(FormInstance).filter(FormInstance.id == form_id).first()
        if not form:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Form not found"
            )
        
        template = db.query(Template).filter(Template.id == form.template_id).first()
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template not found"
            )
        
        # Get form data (from version or current)
        if version_id:
            version = db.query(FormVersion).filter(FormVersion.id == version_id).first()
            if not version:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Version not found"
                )
            form_data = version.data_snapshot
        else:
            form_data_record = db.query(FormData).filter(
                FormData.form_instance_id == form_id
            ).first()
            form_data = form_data_record.data if form_data_record else {}
        
        # Generate filled DOCX
        docx_path = DocumentService._fill_docx(
            template_path=template.original_file_path,
            schema=template.schema,
            data=form_data,
            form_id=form_id,
            version_id=version_id
        )
        
        # Convert to PDF
        pdf_path = DocumentService._convert_to_pdf(docx_path)
        
        # Update version with document paths if version_id provided
        if version_id:
            version.generated_docx_path = docx_path
            version.generated_pdf_path = pdf_path
            db.commit()
        
        return docx_path, pdf_path
    
    @staticmethod
    def _fill_docx(
        template_path: str,
        schema: Dict[str, Any],
        data: Dict[str, Any],
        form_id: int,
        version_id: Optional[int]
    ) -> str:
        """Fill a DOCX template with form data."""
        # Create output directory
        output_dir = os.path.join(settings.generated_dir, str(form_id))
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate output filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        version_suffix = f"_v{version_id}" if version_id else ""
        output_filename = f"form_{form_id}{version_suffix}_{timestamp}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        # Copy template to output
        shutil.copy(template_path, output_path)
        
        # Open and fill document
        doc = Document(output_path)
        fields = schema.get("fields", [])

        # Build field lookup by ID
        field_map = {f["id"]: f for f in fields}

        # Flatten nested data to dot-notation keys (e.g., personnel.pi_name)
        flat_data = DocumentService._flatten_dict(data)

        # Handle Section C contact fields specially (row 9 is fully merged)
        # Combine contact_name, contact_ext, contact_email into one value
        contact_fields = ['personnel.contact_name', 'personnel.contact_ext', 'personnel.contact_email']
        contact_values = {f: flat_data.get(f, '') for f in contact_fields}
        if any(contact_values.values()):
            # Format: "Name | Ext: xxx | Email: xxx"
            parts = []
            if contact_values['personnel.contact_name']:
                parts.append(str(contact_values['personnel.contact_name']))
            if contact_values['personnel.contact_ext']:
                parts.append(f"Ext: {contact_values['personnel.contact_ext']}")
            if contact_values['personnel.contact_email']:
                parts.append(f"Email: {contact_values['personnel.contact_email']}")
            combined_contact = " | ".join(parts)

            # Write to the merged row 9 cell
            if len(doc.tables) > 1:
                table = doc.tables[1]
                if len(table.rows) > 9:
                    cell = table.rows[9].cells[0]
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        para.clear()
                        para.add_run(combined_contact)
                    else:
                        cell.text = combined_contact

        # Handle Funding fields specially (Table 0, Row 14)
        # The cell has multiple paragraphs - need to modify each one individually
        funding_intramural = flat_data.get('funding.intramural_dept', '')
        funding_extramural = flat_data.get('funding.extramural_sponsor', '')
        funding_federal = flat_data.get('funding.federally_funded', '')

        if funding_intramural or funding_extramural or funding_federal:
            if len(doc.tables) > 0:
                table = doc.tables[0]
                if len(table.rows) > 14:
                    cell = table.rows[14].cells[0]

                    # Modify each paragraph individually
                    for para in cell.paragraphs:
                        para_text = para.text

                        if funding_intramural and 'department or fund' in para_text.lower():
                            # Clear and rewrite this paragraph
                            para.clear()
                            para.add_run(f"If intramural, what department or fund? {funding_intramural}")

                        elif funding_extramural and 'name of the sponsor' in para_text.lower():
                            para.clear()
                            para.add_run(f"If extramural, what is the name of the sponsor? {funding_extramural}")

                        elif funding_federal and 'federally funded' in para_text.lower():
                            federal_label = "Yes" if funding_federal == "yes" else "No"
                            para.clear()
                            para.add_run(f"Is the study federally funded? {federal_label}")

        # Handle Section V.B - Drugs/Biologics/Devices fields
        drugs_biologics_devices = flat_data.get('study.drugs_biologics_devices', [])
        fda_regulations = flat_data.get('study.fda_regulations_apply', '')
        ind_number = flat_data.get('study.ind_number', '')
        ide_number = flat_data.get('study.ide_number', '')
        study_phase = flat_data.get('study.phase', [])
        schedule_drugs = flat_data.get('study.schedule_drugs', '')
        schedule_approval = flat_data.get('study.schedule_drugs_approval', '')

        # Fill paragraph 20 - drugs/biologics/devices checkboxes
        if isinstance(drugs_biologics_devices, list) and len(drugs_biologics_devices) > 0:
            if len(doc.paragraphs) > 20:
                para = doc.paragraphs[20]
                text = para.text
                # Mark which items are checked
                if 'drugs' in drugs_biologics_devices:
                    text = text.replace('drugs ,', 'drugs ☑,')
                if 'biologics' in drugs_biologics_devices:
                    text = text.replace('biologics ,', 'biologics ☑,')
                if 'devices' in drugs_biologics_devices:
                    text = text.replace('devices  ', 'devices ☑ ')
                para.clear()
                para.add_run(text)

            # Para 22 - Mark Yes and FDA regulations
            if len(doc.paragraphs) > 22:
                para = doc.paragraphs[22]
                text = para.text
                if fda_regulations == 'yes':
                    text = text.replace('Yes:', 'Yes ☑:')
                para.clear()
                para.add_run(text)
        else:
            # No drugs/biologics/devices - mark No in para 21
            if len(doc.paragraphs) > 21:
                para = doc.paragraphs[21]
                para.clear()
                para.add_run("No ☑")

        # Fill paragraph 23 - IND/IDE numbers (when FDA regulations = Yes)
        if fda_regulations == 'yes':
            if len(doc.paragraphs) > 23:
                para = doc.paragraphs[23]
                text = f"Yes ☑:  Provide IND# {ind_number or '______'}  IDE# {ide_number or '______'}"
                para.clear()
                para.add_run(text)
        elif fda_regulations == 'no' and len(doc.paragraphs) > 24:
            # Mark No for FDA regulations in paragraph 24
            para = doc.paragraphs[24]
            text = para.text
            if text.startswith('No:'):
                text = text.replace('No:', 'No ☑:', 1)
                para.clear()
                para.add_run(text)

        # Fill paragraph 25 - Phase checkboxes
        phase_other_text = flat_data.get('study.phase_other', '')
        if isinstance(study_phase, list) and len(study_phase) > 0:
            if len(doc.paragraphs) > 25:
                para = doc.paragraphs[25]
                text = "2.  Check appropriately:"
                for p in ['phase1', 'phase2', 'phase3', 'phase4', 'emergency']:
                    label = {'phase1': 'Phase I', 'phase2': 'Phase II', 'phase3': 'Phase III', 'phase4': 'Phase IV', 'emergency': 'Emergency Use'}[p]
                    if p in study_phase:
                        text += f" {label} ☑,"
                    else:
                        text += f" {label} ,"
                # Handle Other with specify field
                if 'other' in study_phase:
                    text += f" Other ☑, specify: {phase_other_text}"
                else:
                    text += " Other, specify:"
                para.clear()
                para.add_run(text)

        # Fill paragraph 27-28 - Schedule drugs Yes/No
        if schedule_drugs == 'no' and len(doc.paragraphs) > 27:
            para = doc.paragraphs[27]
            para.clear()
            para.add_run("No ☑")
        elif schedule_drugs == 'yes' and len(doc.paragraphs) > 28:
            # Mark Yes and fill approval info in paragraph 28
            para = doc.paragraphs[28]
            para.clear()
            approval_text = "Yes ☑: This use must be approved by the California State Research Advisory Panel.  Indicate whether "
            if schedule_approval == 'investigator':
                approval_text += "you ☑ or the sponsor  will obtain this approval."
            elif schedule_approval == 'sponsor':
                approval_text += "you  or the sponsor ☑ will obtain this approval."
            else:
                approval_text += "you  or the sponsor  will obtain this approval."
            para.add_run(approval_text)

        # Handle Section V.A - Study Initiation (paragraphs 14-18)
        study_initiated_by = flat_data.get('study.initiated_by', '')
        study_initiated_other = flat_data.get('study.initiated_other', '')
        if study_initiated_by:
            # Map form values to paragraph text patterns
            initiation_map = {
                'local': 'Local investigator',
                'sponsor_investigator': 'sponsor-investigator',
                'cooperative': 'Cooperative group',
                'external': 'External sponsor',
                'other': 'Other, specify'
            }
            target_text = initiation_map.get(study_initiated_by, '')
            if target_text:
                DocumentService._check_checkbox_by_text(doc, target_text, True)

            # Fill in "Other, specify" text if selected
            if study_initiated_by == 'other' and study_initiated_other:
                for para in doc.paragraphs:
                    if 'other, specify' in para.text.lower():
                        # Add the specified text after the checkbox
                        for run in para.runs:
                            if '\u2002' in run.text or '     ' in run.text:  # En-space placeholder
                                run.text = study_initiated_other
                                break
                        break

        # Handle Section V.C - Student Project (paragraph 32)
        is_student_project = flat_data.get('study.is_student_project', '')
        student_name = flat_data.get('study.student_name', '')
        student_program = flat_data.get('study.student_program', '')
        if is_student_project:
            # Find paragraph 32 which has " No  Yes" pattern
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text == 'No  Yes' or para_text == 'No Yes':
                    # Check previous paragraph to confirm it's student project
                    if i > 0 and 'student project' in doc.paragraphs[i-1].text.lower():
                        # This paragraph has two checkboxes: No then Yes
                        checkbox_count = 0
                        for run in para.runs:
                            ffData = run._r.find('.//' + qn('w:ffData'))
                            if ffData is not None:
                                checkbox = ffData.find(qn('w:checkBox'))
                                if checkbox is not None:
                                    # Remove existing checked
                                    existing = checkbox.find(qn('w:checked'))
                                    if existing is not None:
                                        checkbox.remove(existing)
                                    # Check appropriate one (0=No, 1=Yes)
                                    should_check = (checkbox_count == 0 and is_student_project == 'no') or \
                                                   (checkbox_count == 1 and is_student_project == 'yes')
                                    if should_check:
                                        checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                        checked_elem.set(qn('w:val'), '1')
                                    checkbox_count += 1
                        break

        # Handle Section V.D - Ionizing Radiation (nested checkboxes in paragraphs 35-45)
        # Structure:
        # Para 35: No (main) / Para 36: Yes (main)
        # Para 38: No (direct benefit) / Para 39: Yes (direct benefit)
        # Para 41: No (routine) / Para 42: Yes (routine)
        # Para 44: No (greater dose) / Para 45: Yes (greater dose)
        ionizing_radiation = flat_data.get('study.ionizing_radiation', '')
        radiation_direct_benefit = flat_data.get('study.radiation_direct_benefit', '')
        radiation_routine = flat_data.get('study.radiation_routine', '')
        radiation_greater_dose = flat_data.get('study.radiation_greater_dose', '')

        if ionizing_radiation:
            # Main question: Para 35 (No) or Para 36 (Yes)
            if ionizing_radiation == 'no':
                # Para 35: "No: Radiation Safety Committee (RSC) review not required"
                DocumentService._check_legacy_checkbox(doc.paragraphs[35], True)
            else:
                # Para 36: "Yes: Will participants in this study receive direct medical benefits?"
                DocumentService._check_legacy_checkbox(doc.paragraphs[36], True)

        if ionizing_radiation == 'yes' and radiation_direct_benefit:
            # Direct benefit question: Para 38 (No) or Para 39 (Yes)
            if radiation_direct_benefit == 'no':
                # Para 38: "No: RSC review REQUIRED"
                DocumentService._check_legacy_checkbox(doc.paragraphs[38], True)
            else:
                # Para 39: "Yes: Is the proposed use..."
                DocumentService._check_legacy_checkbox(doc.paragraphs[39], True)

        if ionizing_radiation == 'yes' and radiation_direct_benefit == 'yes' and radiation_routine:
            # Routine question: Para 41 (No) or Para 42 (Yes)
            if radiation_routine == 'no':
                # Para 41: "No: RSC review REQUIRED"
                DocumentService._check_legacy_checkbox(doc.paragraphs[41], True)
            else:
                # Para 42: "Yes: Will subjects participating..."
                DocumentService._check_legacy_checkbox(doc.paragraphs[42], True)

        if ionizing_radiation == 'yes' and radiation_direct_benefit == 'yes' and radiation_routine == 'yes' and radiation_greater_dose:
            # Greater dose question: Para 44 (No) or Para 45 (Yes)
            if radiation_greater_dose == 'no':
                # Para 44: "No: RSC review not required"
                DocumentService._check_legacy_checkbox(doc.paragraphs[44], True)
            else:
                # Para 45: "Yes: RSC review REQUIRED"
                DocumentService._check_legacy_checkbox(doc.paragraphs[45], True)

        # Handle Section V.E - SCRO (paragraph 48)
        scro_required = flat_data.get('study.scro_required', '')
        if scro_required:
            # Paragraph 48 has "No     Yes: Letter of approval from SCRO"
            for i, para in enumerate(doc.paragraphs):
                if 'scro' in para.text.lower() and ('no' in para.text.lower() and 'yes' in para.text.lower()):
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and scro_required == 'no') or \
                                               (checkbox_count == 1 and scro_required == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # Handle Section V.F - IBC (paragraphs 53, 56, 57)
        ibc_infectious = flat_data.get('study.ibc_infectious', '')
        ibc_recombinant = flat_data.get('study.ibc_recombinant', '')
        ibc_hazardous = flat_data.get('study.ibc_hazardous', '')

        # IBC Infectious agents (para 53: "No     Yes: Letter of approval from IBC")
        if ibc_infectious:
            for para in doc.paragraphs:
                if 'letter of approval from ibc' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and ibc_infectious == 'no') or \
                                               (checkbox_count == 1 and ibc_infectious == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # IBC Recombinant (para 56: "No     Yes: Recombinant or synthetic")
        if ibc_recombinant:
            for para in doc.paragraphs:
                if 'recombinant or synthetic' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and ibc_recombinant == 'no') or \
                                               (checkbox_count == 1 and ibc_recombinant == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # IBC Hazardous (para 57: "No    Yes: Carcinogens, mutagens")
        if ibc_hazardous:
            for para in doc.paragraphs:
                if 'carcinogens' in para.text.lower() or 'hazardous materials' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and ibc_hazardous == 'no') or \
                                               (checkbox_count == 1 and ibc_hazardous == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # Handle Section V.G - PHS Submission (paragraph 60)
        phs_submission = flat_data.get('study.phs_submission', '')
        if phs_submission:
            for para in doc.paragraphs:
                if 'phs policy' in para.text.lower() or 'public health service' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and phs_submission == 'no') or \
                                               (checkbox_count == 1 and phs_submission == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # Handle Section VI - Classification of Subjects (checkboxes in Table 2, Row 3)
        vulnerable_populations = flat_data.get('population.vulnerable', [])
        special_populations = flat_data.get('population.special', [])
        special_other = flat_data.get('population.special_other', '')
        other_populations = flat_data.get('population.other_populations', [])
        other_populations_other = flat_data.get('population.other_populations_other', '')

        # Classification is in Table 2, Row 3
        if len(doc.tables) > 2:
            table = doc.tables[2]
            if len(table.rows) > 3:
                row = table.rows[3]

                # Cell 0 = Vulnerable populations
                if len(row.cells) > 0 and isinstance(vulnerable_populations, list):
                    cell = row.cells[0]
                    pop_labels = {
                        'abortuses': 'abortuses',
                        'diminished_capacity': 'diminished',
                        'economically_disadvantaged': 'economically disadvantaged',
                        'educationally_disadvantaged': 'educationally disadvantaged',
                        'fetuses': 'fetuses',
                        'minors': 'minors',
                        'neonates': 'neonates',
                        'pregnant': 'pregnant',
                        'prisoners': 'prisoners'
                    }
                    for para in cell.paragraphs:
                        para_text_lower = para.text.lower()
                        for pop_value, pop_search in pop_labels.items():
                            if pop_search in para_text_lower and pop_value in vulnerable_populations:
                                DocumentService._check_legacy_checkbox(para, True)
                                break

                # Cell 1 = Special populations
                if len(row.cells) > 1 and isinstance(special_populations, list):
                    cell = row.cells[1]
                    pop_labels = {
                        'elderly': 'elderly',
                        'illiterate': 'illiterate',
                        'institutionalized': 'institutionalized',
                        'patients_inpatients': 'inpatients',
                        'patients_outpatients': 'outpatients',
                        'terminally_ill': 'terminally',
                        'traumatized': 'traumatized'
                    }
                    for para in cell.paragraphs:
                        para_text_lower = para.text.lower()
                        for pop_value, pop_search in pop_labels.items():
                            if pop_search in para_text_lower and pop_value in special_populations:
                                DocumentService._check_legacy_checkbox(para, True)
                                break
                        # Handle "Other, specify" for special populations
                        if special_other and 'other, specify' in para_text_lower:
                            DocumentService._check_legacy_checkbox(para, True)
                            # Add the specified text
                            for run in para.runs:
                                if '\u2002' in run.text:  # En-space placeholder
                                    run.text = special_other
                                    break

                # Cell 2 = Other populations
                if len(row.cells) > 2 and isinstance(other_populations, list):
                    cell = row.cells[2]
                    pop_labels = {
                        'employees': 'employees',
                        'female_only': 'female (excludes males)',
                        'healthy': 'healthy (non-patient)',
                        'male_only': 'male (excludes females)',
                        'minorities': 'minorities',
                        'non_english': 'non-english',
                        'physically_handicapped': 'physically handicapped',
                        'sda_cohort': 'seventh-day adventist',
                        'students': 'students'
                    }
                    for para in cell.paragraphs:
                        para_text_lower = para.text.lower()
                        for pop_value, pop_search in pop_labels.items():
                            if pop_search in para_text_lower and pop_value in other_populations:
                                DocumentService._check_legacy_checkbox(para, True)
                                break
                        # Handle "Other, specify" for other populations
                        if other_populations_other and 'other, specify' in para_text_lower:
                            DocumentService._check_legacy_checkbox(para, True)
                            for run in para.runs:
                                if '\u2002' in run.text:
                                    run.text = other_populations_other
                                    break

        # Handle Section VI - Recruitment fields
        recruitment_source = flat_data.get('recruitment.source', [])
        recruitment_source_other = flat_data.get('recruitment.source_other', '')
        flyers_used = flat_data.get('recruitment.flyers_used', '')
        verbal_used = flat_data.get('recruitment.verbal_used', '')
        electronic_used = flat_data.get('recruitment.electronic_used', '')
        electronic_description = flat_data.get('recruitment.electronic_description', '')

        if isinstance(recruitment_source, list):
            # Para 72: own_patients - "PI/collaborators will recruit"
            if 'own_patients' in recruitment_source:
                DocumentService._check_checkbox_by_text(doc, 'PI/collaborators will recruit', True)
            # Para 73: letter_referrals - "letter to colleagues asking for referrals"
            if 'letter_referrals' in recruitment_source:
                DocumentService._check_checkbox_by_text(doc, 'asking for referrals', True)
            # Para 74: letter_patients - "Dear Patient"
            if 'letter_patients' in recruitment_source:
                DocumentService._check_checkbox_by_text(doc, 'Dear Patient', True)
            # Para 75: other - fill text if selected
            if 'other' in recruitment_source and recruitment_source_other:
                for para in doc.paragraphs:
                    if 'other, specify' in para.text.lower() and 'source' not in para.text.lower():
                        # Add the text to the specify field
                        for run in para.runs:
                            if '\u2002' in run.text or '     ' in run.text:
                                run.text = recruitment_source_other
                                break
                        break

        # Para 77: Flyers Yes/No (has both checkboxes)
        if flyers_used:
            for para in doc.paragraphs:
                if 'flyers' in para.text.lower() and 'attach copy' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and flyers_used == 'no') or \
                                               (checkbox_count == 1 and flyers_used == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # Para 79: Verbal recruitment Yes/No
        if verbal_used:
            for para in doc.paragraphs:
                if 'verbal' in para.text.lower() and 'script' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and verbal_used == 'no') or \
                                               (checkbox_count == 1 and verbal_used == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # Para 81: Electronic recruitment Yes/No
        if electronic_used:
            for para in doc.paragraphs:
                if 'yes, describe' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                should_check = (checkbox_count == 0 and electronic_used == 'no') or \
                                               (checkbox_count == 1 and electronic_used == 'yes')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                        # Fill description if yes
                        if electronic_used == 'yes' and electronic_description:
                            if '\u2002' in run.text or '     ' in run.text:
                                run.text = electronic_description
                    break

        # Handle Section VI - Consent fields
        consent_timing = flat_data.get('consent.timing', '')
        consent_documents = flat_data.get('consent.documents_required', [])
        consent_waiver = flat_data.get('consent.waiver_requested', [])
        consent_inducement = flat_data.get('consent.inducement', '')

        # Para 84: Consent timing (has 2 checkboxes: In conjunction / At a separate appointment)
        if consent_timing:
            for para in doc.paragraphs:
                if 'relative to the performance' in para.text.lower():
                    checkbox_count = 0
                    for run in para.runs:
                        ffData = run._r.find('.//' + qn('w:ffData'))
                        if ffData is not None:
                            checkbox = ffData.find(qn('w:checkBox'))
                            if checkbox is not None:
                                existing = checkbox.find(qn('w:checked'))
                                if existing is not None:
                                    checkbox.remove(existing)
                                # First checkbox = conjunction, second = separate
                                should_check = (checkbox_count == 0 and consent_timing == 'conjunction') or \
                                               (checkbox_count == 1 and consent_timing == 'separate')
                                if should_check:
                                    checked_elem = etree.SubElement(checkbox, qn('w:checked'))
                                    checked_elem.set(qn('w:val'), '1')
                                checkbox_count += 1
                    break

        # Consent documents checkboxes (Para 86-90)
        if isinstance(consent_documents, list):
            doc_labels = {
                'informed_consent': 'informed consent document',
                'parent_permission': 'consent/permission of parent',
                'assent_13_17': '13 – 17 yrs',
                'assent_7_12': '7 – 12 yrs',
                'hipaa_auth': 'authorization for use of protected health'
            }
            for doc_type, search_text in doc_labels.items():
                if doc_type in consent_documents:
                    DocumentService._check_checkbox_by_text(doc, search_text, True)

        # Consent waiver checkboxes (Para 92-94)
        if isinstance(consent_waiver, list):
            waiver_labels = {
                'waiver_consent': 'waiver of consent',
                'waiver_written': 'waiver of written consent',
                'waiver_signed': 'waiver of signed consent',
                'waiver_hipaa': 'waiver of hipaa'
            }
            for waiver_type, search_text in waiver_labels.items():
                if waiver_type in consent_waiver:
                    DocumentService._check_checkbox_by_text(doc, search_text, True)

        # Inducement text
        if consent_inducement:
            for para in doc.paragraphs:
                if 'inducement' in para.text.lower() and 'offered' in para.text.lower():
                    # Find text field placeholder and fill it
                    for run in para.runs:
                        if '\u2002' in run.text or '     ' in run.text:
                            run.text = consent_inducement
                            break
                    break

        # Fill each field
        for field_id, value in flat_data.items():
            # Skip contact fields as they were handled above
            if field_id in contact_fields:
                continue
            if field_id.startswith("_"):
                continue  # Skip system fields

            field_def = field_map.get(field_id)
            if not field_def:
                continue

            field_type = field_def.get("type", "text")

            # Handle checkbox fields specially
            if field_type == "checkbox" and isinstance(value, list):
                DocumentService._fill_checkbox_field(doc, value, field_def)
                continue

            anchor = field_def.get("anchor")
            if not anchor:
                continue

            DocumentService._write_value_to_anchor(doc, anchor, value, field_def)

        # Save the filled document
        doc.save(output_path)
        
        return output_path

    @staticmethod
    def _fix_header_table_borders(doc: Document) -> None:
        """
        Fix header table (Table 0) borders for LibreOffice compatibility.
        Cell 0 (logo area) should have no visible borders.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        if len(doc.tables) == 0:
            return

        table = doc.tables[0]
        if len(table.rows) == 0:
            return

        # Fix Cell 0 (logo cell) - remove all borders
        cell = table.rows[0].cells[0]
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        # Remove existing borders
        existing_borders = tc_pr.find(qn('w:tcBorders'))
        if existing_borders is not None:
            tc_pr.remove(existing_borders)

        # Add new borders element with all sides set to nil
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tc_borders.append(border)
        tc_pr.append(tc_borders)

    @staticmethod
    def _fill_checkbox_field(
        doc: Document,
        selected_values: list,
        field_def: Dict[str, Any]
    ) -> None:
        """Fill checkbox fields by updating form field checkboxes or adding visual markers."""
        from docx.oxml.ns import qn

        options = field_def.get("options", [])

        # Build a map of option values to their labels
        option_labels = {opt.get("value"): opt.get("label", opt.get("value")) for opt in options}

        # Find paragraphs containing checkbox options and update them
        for para in doc.paragraphs:
            para_text = para.text.strip()

            # Check each option
            for opt_value, opt_label in option_labels.items():
                # Check if this paragraph contains this option's label
                if opt_label.lower() in para_text.lower():
                    is_selected = opt_value in selected_values

                    # Try to find and update form field checkboxes
                    for run in para.runs:
                        run_xml = run._r.xml
                        if 'fldChar' in run_xml or 'checkBox' in run_xml.lower():
                            # Found a form field - try to update it
                            for elem in run._r.iter():
                                # Look for checkbox default value
                                if elem.tag.endswith('default') or elem.tag.endswith('checked'):
                                    elem.set(qn('w:val'), '1' if is_selected else '0')

                    # Also try to find and replace checkbox symbols
                    # Common unchecked: ☐ (U+2610), □ (U+25A1)
                    # Common checked: ☑ (U+2611), ☒ (U+2612), ■ (U+25A0)
                    for run in para.runs:
                        if '☐' in run.text:
                            run.text = run.text.replace('☐', '☑' if is_selected else '☐')
                        elif '□' in run.text:
                            run.text = run.text.replace('□', '■' if is_selected else '□')

                    break  # Found the option, move to next

    @staticmethod
    def _write_value_to_anchor(
        doc: Document,
        anchor: Dict[str, Any],
        value: Any,
        field_def: Dict[str, Any]
    ) -> None:
        """Write a value to the anchored location in the document."""
        anchor_type = anchor.get("type")
        
        if anchor_type == "paragraph":
            DocumentService._fill_paragraph_anchor(doc, anchor, value, field_def)
        elif anchor_type == "table_cell":
            DocumentService._fill_table_cell_anchor(doc, anchor, value, field_def)
        elif anchor_type == "table":
            DocumentService._fill_table_anchor(doc, anchor, value, field_def)
    
    @staticmethod
    def _fill_paragraph_anchor(
        doc: Document,
        anchor: Dict[str, Any],
        value: Any,
        field_def: Dict[str, Any]
    ) -> None:
        """Fill a paragraph-anchored field."""
        paragraph_contains = anchor.get("paragraph_contains", "")
        paragraph_index = anchor.get("paragraph_index")
        
        for i, para in enumerate(doc.paragraphs):
            # Match by index or by content
            if paragraph_index is not None and i == paragraph_index:
                DocumentService._insert_value_in_paragraph(para, value, field_def)
                return
            elif paragraph_contains and paragraph_contains.lower() in para.text.lower():
                DocumentService._insert_value_in_paragraph(para, value, field_def)
                return
    
    @staticmethod
    def _insert_value_in_paragraph(
        para: Paragraph,
        value: Any,
        field_def: Dict[str, Any]
    ) -> None:
        """Insert value into a paragraph, preserving formatting."""
        text = para.text
        field_type = field_def.get("type", "text")
        
        # Format value based on type
        if field_type == "checkbox":
            if isinstance(value, list):
                formatted_value = ", ".join(str(v) for v in value)
            else:
                formatted_value = "☑" if value else "☐"
        elif isinstance(value, list):
            formatted_value = ", ".join(str(v) for v in value)
        else:
            formatted_value = str(value) if value else ""
        
        # Find placeholder pattern (underscores, colon, question mark, or append)
        patterns = [
            (r'(_+)', formatted_value),  # Replace underscores
            (r'(:\s*)$', f': {formatted_value}'),  # Append after colon at end
            (r'(\?\s*)$', f'? {formatted_value}'),  # Append after question mark at end
        ]

        matched = False
        new_text = text
        for pattern, replacement in patterns:
            if re.search(pattern, text):
                new_text = re.sub(pattern, replacement, text, count=1)
                matched = True
                break

        # If no pattern matched, append the value at the end
        if not matched and formatted_value:
            new_text = f"{text.rstrip()} {formatted_value}"

        # Split text to separate original label from the filled value
        # so we can make only the value bold
        if matched:
            # Find where the value was inserted
            for pattern, replacement in patterns:
                match = re.search(pattern, text)
                if match:
                    prefix = text[:match.start()]
                    suffix = text[match.end():]
                    break
            else:
                prefix = text
                suffix = ""

            # Clear paragraph and add runs with formatting
            para.clear()
            if prefix:
                # Add space before value if prefix doesn't end with space
                if not prefix.endswith(' '):
                    prefix = prefix + ' '
                para.add_run(prefix)
            value_run = para.add_run(formatted_value)
            value_run.underline = True
            if suffix:
                para.add_run(suffix)
        else:
            # For appended values
            para.clear()
            para.add_run(text.rstrip() + "  ")  # Two spaces before value
            value_run = para.add_run(formatted_value)
            value_run.underline = True
        return
    
    @staticmethod
    def _fill_table_cell_anchor(
        doc: Document,
        anchor: Dict[str, Any],
        value: Any,
        field_def: Dict[str, Any]
    ) -> None:
        """Fill a table cell-anchored field."""
        table_index = anchor.get("table_index", 0)
        row_index = anchor.get("row_index", 0)
        column_index = anchor.get("column_index", 1)

        if table_index >= len(doc.tables):
            return

        table = doc.tables[table_index]

        if row_index >= len(table.rows):
            return

        row = table.rows[row_index]

        if column_index >= len(row.cells):
            return

        cell = row.cells[column_index]

        # Format value based on field type
        field_type = field_def.get("type", "text")
        if field_type == "checkbox":
            # For checkbox fields, show Yes/checkmark if checked
            if isinstance(value, list) and len(value) > 0:
                formatted_value = "Yes"
            elif value:
                formatted_value = "Yes"
            else:
                formatted_value = ""
        elif field_type == "radio":
            # For radio fields, show the label of the selected option
            options = field_def.get("options", [])
            option_map = {opt.get("value"): opt.get("label", opt.get("value")) for opt in options}
            formatted_value = option_map.get(value, str(value) if value else "")
        elif isinstance(value, list):
            formatted_value = ", ".join(str(v) for v in value)
        else:
            formatted_value = str(value) if value else ""

        # Get current cell text
        current_text = cell.text.strip() if cell.text else ""

        # Check for special anchor options
        append_to_label = anchor.get("append_to_label", False)
        replace_pattern = anchor.get("replace_pattern")
        replace_with = anchor.get("replace_with")
        cell_paragraph_index = anchor.get("cell_paragraph_index")

        # If specific paragraph index is specified, write to that paragraph
        if cell_paragraph_index is not None and formatted_value:
            if cell_paragraph_index < len(cell.paragraphs):
                para = cell.paragraphs[cell_paragraph_index]
                para.clear()
                para.add_run(formatted_value)
            return

        if replace_pattern and replace_with and formatted_value:
            # Replace pattern in existing text
            import re
            new_text = re.sub(
                replace_pattern,
                replace_with.replace("{value}", formatted_value),
                current_text
            )
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
                para.add_run(new_text)
            else:
                cell.text = new_text
        elif append_to_label and current_text and formatted_value:
            # Append value to existing label text
            new_text = f"{current_text} {formatted_value}"
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
                para.add_run(new_text)
            else:
                cell.text = new_text
        else:
            # Set cell text (replace)
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
                para.add_run(formatted_value)
            else:
                cell.text = formatted_value
    
    @staticmethod
    def _fill_table_anchor(
        doc: Document,
        anchor: Dict[str, Any],
        value: Any,
        field_def: Dict[str, Any]
    ) -> None:
        """Fill a repeatable table-anchored field."""
        if not isinstance(value, list):
            return

        table_index = anchor.get("table_index", 0)
        start_row = anchor.get("start_row", 1)  # Default to row 1 if not specified

        if table_index >= len(doc.tables):
            return

        table = doc.tables[table_index]
        repeatable_config = field_def.get("repeatable_config", {})
        columns = repeatable_config.get("columns", [])

        # Get column mapping - accounts for merged cells
        # Map logical column index to actual cell index
        column_mapping = repeatable_config.get("column_mapping", None)

        # Fill data rows starting from start_row
        for row_idx, row_data in enumerate(value):
            actual_row_idx = start_row + row_idx
            if actual_row_idx >= len(table.rows):
                # Need to add new row
                table.add_row()

            row = table.rows[actual_row_idx]

            for col_idx, col_def in enumerate(columns):
                # Use column_mapping if provided, otherwise use col_idx
                if column_mapping and col_idx < len(column_mapping):
                    actual_col_idx = column_mapping[col_idx]
                else:
                    actual_col_idx = col_idx

                if actual_col_idx >= len(row.cells):
                    continue

                cell = row.cells[actual_col_idx]
                col_id = col_def.get("id", f"col_{col_idx}")
                cell_value = row_data.get(col_id, "") if isinstance(row_data, dict) else ""

                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    para.clear()
                    para.add_run(str(cell_value))
                else:
                    cell.text = str(cell_value)
    
    @staticmethod
    def _convert_to_pdf(docx_path: str) -> str:
        """Convert DOCX to PDF using LibreOffice."""
        output_dir = os.path.dirname(docx_path)
        
        # LibreOffice command for headless PDF conversion
        cmd = [
            settings.libreoffice_path,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60  # 60 second timeout
            )
            
            if result.returncode != 0:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"PDF conversion failed: {result.stderr}"
                )
            
            # PDF file has same name with .pdf extension
            pdf_path = docx_path.replace('.docx', '.pdf')
            
            if not os.path.exists(pdf_path):
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF file was not created"
                )
            
            return pdf_path
            
        except subprocess.TimeoutExpired:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="PDF conversion timed out"
            )
        except FileNotFoundError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="LibreOffice not found. Please install LibreOffice."
            )
    
    @staticmethod
    def get_document_paths(
        db: Session,
        form_id: int,
        version_id: int
    ) -> Optional[Tuple[str, str]]:
        """Get existing document paths for a version."""
        version = db.query(FormVersion).filter(
            FormVersion.id == version_id
        ).first()
        
        if not version:
            return None
        
        if version.generated_docx_path and version.generated_pdf_path:
            if os.path.exists(version.generated_docx_path) and os.path.exists(version.generated_pdf_path):
                return (version.generated_docx_path, version.generated_pdf_path)
        
        return None
